import argparse
import re
from docx import Document
from docx.shared import Pt
import sys

def censor_text(text, censor_list, censor_map):
    for word in censor_list:
        if word not in censor_map:
            censor_map[word] = f'<CENSORED_WORD{len(censor_map) + 1}>'
        text = re.sub(rf'\b{re.escape(word)}\b', censor_map[word], text, flags=re.IGNORECASE)
    return text

def censor_ips_md(file_path, censor_map):
    # Define the regex pattern to match IP addresses
    ip_pattern = re.compile(r'\b(?:\d{1,3}\.){3}\d{1,3}\b')

    # Read the content of the file
    with open(file_path, 'r') as file:
        file_content = file.read()

    # Use regex to find all IP addresses in the content
    ip_addresses = re.findall(ip_pattern, file_content)

    # Replace each IP address with <CENSORED_IPX>
    for i, ip_address in enumerate(ip_addresses):
        censored_ip = f'<CENSORED_IP{censor_map + 1}>'
        file_content = file_content.replace(ip_address, censored_ip)

    # Write the modified content back to the file
    with open(file_path, 'w') as file:
        file.write(file_content)
def censor_ips(text, censor_map):
    ip_pattern = re.compile(r'\b(?:\d{1,3}\.){3}\d{1,3}\b')

    def censor_ip(match):
        ip = match.group()
        if ip not in censor_map:
            censor_map[ip] = f'<CENSOR_IP{len(censor_map) + 1}>'
        return censor_map[ip]

    return ip_pattern.sub(censor_ip, text)


def process_docx(file_path, censor_list, censor_ips_flag, decensor_flag=False, decensor_map=None):
    document = Document(file_path)
    censor_map = {}

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(12)
            if decensor_flag:
                run.text = replace_censored_words(run.text, decensor_map)
            elif censor_ips_flag:
                run.text = censor_ips(run.text, censor_map)
            else:
                pass
            run.text = censor_text(run.text, censor_list, censor_map)

    output_file = 'decensored_output.docx' if decensor_flag else 'censored_output.docx'
    censored_document = Document()

    for paragraph in document.paragraphs:
        new_paragraph = censored_document.add_paragraph()
        new_paragraph.style = paragraph.style
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.size = Pt(12)

    censored_document.save(output_file)

    if not decensor_flag:
        with open('censor_map.txt', 'w') as map_file:
            for key, value in censor_map.items():
                map_file.write(f'{value}:{key}\n')

def process_md(file_path, censor_list, censor_ips_flag, decensor_flag=False, decensor_map=None):
    with open(file_path, 'r') as md_file:
        md_text = md_file.read()

    censor_map = {}
    censored_md_text = censor_text(md_text, censor_list, censor_map)
    if decensor_flag:
        censored_md_text = replace_censored_words(md_text, decensor_map)
    elif censor_ips_flag:
        censored_md_text = censor_ips(censored_md_text, censor_map)

    output_file = 'decensored_output.md' if decensor_flag else 'censored_output.md'
    with open(output_file, 'w') as output_file:
        output_file.write(censored_md_text)

    if not decensor_flag:
        with open('censor_map.txt', 'w') as map_file:
            for key, value in censor_map.items():
                map_file.write(f'{value}:{key}\n')
    # print(censored_md_text)

def replace_censored_words(text, censor_map):
    for censored_word, replacement in censor_map.items():
        text = text.replace(censored_word, replacement)
    return text

def main():
    parser = argparse.ArgumentParser(description='Censor words and/or IP addresses, and create a map.')
    parser.add_argument('file_path', type=str, help='Path to the input file')
    parser.add_argument('--censor', nargs='+', default=[], help='List of words to censor')
    parser.add_argument('--ips', action='store_true', help='Censor all IP addresses')
    parser.add_argument('--decensor', action='store_true', help='Decensor the text using the provided map')
    parser.add_argument('--decensor_map',default='censor_map.txt',type=str, help='Path to the decensor map file')
    args = parser.parse_args()

    file_extension = args.file_path.split('.')[-1].lower()

    if args.decensor and not args.decensor_map:
        print('Error: --decensor_map is required when using --decensor flag.')
        sys.exit(1)

    if args.decensor:
        try:
            with open(args.decensor_map, 'r') as decensor_map_file:
                decensor_map = dict(line.strip().split(':') for line in decensor_map_file)
        except Exception as e:
            print(f"Error reading the decensor map file: {e}")
            sys.exit(1)

    if file_extension == 'docx':
        process_docx(args.file_path, args.censor, args.ips, args.decensor, decensor_map if args.decensor else None)
    elif file_extension == 'md':
        process_md(args.file_path, args.censor, args.ips, args.decensor, decensor_map if args.decensor else None)
    else:
        print('Unsupported file format. Please use DOCX or MD files.')

if __name__ == "__main__":
    main()
