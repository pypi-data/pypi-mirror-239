import io
import os
import re
from datetime import datetime
from zipfile import BadZipFile

import docx
import fitz
import olefile
import textract
from asposewordscloud import WordsApi
from asposewordscloud.models.requests import ConvertDocumentRequest
from google.cloud import vision
from pypdf import PdfReader
from pypdf.errors import PdfStreamError

from .base import BaseService
from .google_drive import DriveService


class TextExtractService(BaseService):
    def __init__(
        self,
        vision_client: vision.ImageAnnotatorClient,
        words_client: WordsApi
    ) -> None:
        self.vision_client = vision_client
        self.words_client = words_client
        self.drive_service = DriveService()
        super().__init__(log_name='vision.service')

    def extract_text_from_file(self, file_bytes, mime_type, detection_type):
        self.logger.log_text(f"Looking for text in document")
        feature = vision.Feature(type_=detection_type)
        input_config = vision.InputConfig(
            content=file_bytes, mime_type=mime_type
        )
        reader = PdfReader(io.BytesIO(file_bytes))
        all_pages = [i + 1 for i in range(0, len(reader.pages))]
        text = ""
        for i in range(0, len(reader.pages), 5):
            end_page = i + 5 if i + 5 < len(reader.pages) else len(reader.pages)
            self.logger.log_text(f"Processing pages {i + 1} to {end_page} of {len(reader.pages)}")
            # noinspection PyTypeChecker
            annotate_request = vision.AnnotateFileRequest(
                input_config=input_config,
                features=[feature],
                pages=all_pages[i:i + 5]
            )
            # noinspection PyTypeChecker
            batch_annotate_request = vision.BatchAnnotateFilesRequest(
                requests=[annotate_request]
            )
            response = self.vision_client.batch_annotate_files(
                request=batch_annotate_request
            )
            text += "".join([r2.full_text_annotation.text for r1 in response.responses for r2 in r1.responses])
            self.logger.log_text(f"Extracted {len(text)} chars from document")
            if len(text) > 65500:
                text = text[0:65500]
                self.logger.log_text(f"Truncating text to 65500 characters")
                break
        return text

    def extract_text_from_image(self, file_bytes):
        self.logger.log_text(f"Looking for text in image")
        image = vision.Image(content=file_bytes)
        response = self.vision_client.text_detection(image=image)
        texts = response.text_annotations
        text = "".join([t.description for t in texts])
        self.logger.log_text(f"Extracted {len(text)} chars from image")
        return text
    
    def get_text_from_file_content(self, file_bytes: bytes, file_ending: str):
        if len(file_bytes) == 0:
            file_text = 'No Text to Extract'
        elif file_ending in ['xps', 'epub', 'mobi', 'fb2', 'cbz', 'svg']:
            doc = fitz.Document(stream=io.BytesIO(file_bytes))
            file_text = ''
            for page in doc.pages():
                file_text += page.get_text()
        elif file_ending == 'pdf':
            try:
                doc = fitz.Document(stream=io.BytesIO(file_bytes))
                file_text = ''
                for page in doc.pages():
                    file_text += page.get_text()
            except Exception as e:
                print(f"Unable to extract pdf with PyMuPDF: {str(e)}")
                print(f"Attempting cloud vision api")
                try:
                    file_text = self.extract_text_from_file(
                        file_bytes=file_bytes,
                        mime_type='application/pdf',
                        detection_type=vision.Feature.Type.DOCUMENT_TEXT_DETECTION
                    )
                except PdfStreamError:
                    file_text = "Unable to extract PDF file contents. Please verify the document is valid."
        elif file_ending in ['jpg', 'jpeg', 'png']:
            file_text = self.extract_text_from_image(
                file_bytes=file_bytes
            )
        elif file_ending == 'docx':
            try:
                temp_file_path = f"/tmp/resume_{datetime.now()}.docx"
                doc_file = open(temp_file_path, "wb")
                doc_file.write(file_bytes)
                file_text = textract.process(temp_file_path).decode()
                doc_file.close()
                os.remove(temp_file_path)
            except Exception as e:
                print(f"Unable to extract docx with textract: {str(e)}")
                print(f"Attempting docx library")
                try:
                    doc = docx.Document(io.BytesIO(file_bytes))
                    file_text = "\n".join([p.text for p in doc.paragraphs])
                except BadZipFile:
                    file_text = "Unable to extract DOCX file contents. Please verify the document is valid."
                except KeyError as e:
                    if 'no relationship of type' not in str(e):
                        raise e
                    file_text = "Unable to extract DOCX file contents. Please verify the document is valid."
                except ValueError as e:
                    if 'is not a Word file' not in str(e):
                        raise e
                    try:
                        ole = olefile.OleFileIO(io.BytesIO(file_bytes))
                        doc = ole.openstream('WordDocument')
                        file_text = str(doc.read())
                        file_text = re.sub(pattern=r'\\x\S+ *(?=\\)', repl='', string=file_text)
                    except Exception as e:
                        print(f"DOCX Extraction and olefile failed: {str(e)}")
                        file_text = "Unable to extract DOCX file contents. Please verify the document is valid."
        elif file_ending == 'doc':
            # first try olefile
            try:
                ole = olefile.OleFileIO(io.BytesIO(file_bytes))
                doc = ole.openstream('WordDocument')
                file_text = str(doc.read())
                file_text = re.sub(pattern=r'\\x\S+ *(?=\\)', repl='', string=file_text)
            except Exception as e:
                print(f"Unable to parse .doc file with olefile: {str(e)}")
                print("Attempting to convert .doc file to PDF.")
                temp_file_path = f"/tmp/resume_{datetime.now()}.doc"
                doc_file = open(temp_file_path, "wb")
                doc_file.write(file_bytes)
                doc_file.close()

                try:
                    pdf_bytes = self.convert_word_to_pdf(file_name=temp_file_path)

                    doc = fitz.Document(stream=io.BytesIO(pdf_bytes))
                    file_text = ''
                    for page in doc.pages():  # iterate the document pages
                        file_text += page.get_text()
                except Exception as e:
                    print(f"Unable to convert doc to PDF: {str(e)}")
                    print("Attempting to parse .doc file with aspose.")
                    doc_file = open(temp_file_path, "rb")
                    request = ConvertDocumentRequest(document=doc_file, format="docx")
                    docx_content = self.words_client.convert_document(request=request)
                    doc = docx.Document(io.BytesIO(docx_content))
                    file_text = "\n".join([p.text for p in doc.paragraphs])
                    doc_file.close()
                finally:
                    os.remove(temp_file_path)
        elif file_ending == 'xlsx':
            temp_file_path = f"/tmp/resume_{datetime.now()}.xlsx"
            try:
                doc_file = open(temp_file_path, "wb")
                doc_file.write(file_bytes)
                file_text = textract.process(temp_file_path).decode()
                doc_file.close()
            finally:
                os.remove(temp_file_path)
        elif file_ending == 'txt':
            file_text = file_bytes.decode()
        else:
            file_text = 'File Type Not Supported'
        file_text = 'No Text to Extract' if not file_text or len(file_text) == 0 else file_text
        file_text = file_text[0:60000]
        return re.sub(r'[Ss]cript', '_script', file_text)\
            .replace('<', '(<)')\
            .replace('>', '(>)')\
            .replace(r'etc/hosts', 'etc_hosts')\
            .replace('.config', '_config')\
            .replace('/etc/', '_etc_')\
            .replace('/var/log/wtmp', '_var_log_wtmp')\
            .replace('Web.Config', 'Web_Config')

    def convert_word_to_pdf(self, file_name: str):
        file_id = None
        try:
            file_id = self.drive_service.upload_google_doc(file_name=file_name, mime_type='application/msword')
            return self.drive_service.export_pdf(file_id=file_id)
        finally:
            if file_id:
                self.drive_service.delete_file(file_id)
