def equa (poserrorratedata,negerrorratedata,csv_stats,Criterion,doc_loc,a,b,equation,limita,limitb):

    import os,glob, pandas as pd
    import numpy as np  
    import pandas as pd  
    import matplotlib.pyplot as plt
    import seaborn as sns
    import numexpr as ne
    import docx 
    from docx.shared import Inches
    from sklearn.metrics import mean_absolute_error, mean_squared_error, r2_score
    from pymongo import MongoClient



    
    #Generate into Document
    document = docx.Document()
 
    
    heading_text = f"Report of {Criterion} Lunar Crescent Visibility Criterion"
    document.add_heading(heading_text)
    
    document.add_heading('General Statistics of Criterion Variable', level=2)
    aparameter_text = f"General Statistics for  {a} Parameter"
    document.add_heading(aparameter_text, level=3)
    
    #document.save(doc_loc)

    #path = 'C:/Users/syazw/OneDrive - Universiti Malaya/PhD/References/Source of Moon Sighting Report/Raw Data/Final Data/Final.csv'
    #df = pd.read_csv(path, index_col=0)
    
    client = MongoClient("mongodb+srv://userhilalpy:kSTkJBzJ8CbCATmm@hilalpy.wdazfdi.mongodb.net/?retryWrites=true&w=majority")
    db = client.test
    db = client['datahillalpy']
    collection = db['lunarcrescentdatahilalpy']

    # read data from MongoDB collection
    data_mongo = []
    for doc in collection.find():
        data_mongo.append(doc)

# convert data to Pandas DataFrame
    df = pd.DataFrame(data_mongo)
    df[a] = abs(df[a])
    #df=df[(df[x] <= 20)]

    #Only take Visible Data
    df = df[df['V'] =='V']

    #Only Take Evening Observation
    df = df[df['O'] =='E']

    ##Generate the Boxplot
    chart =sns.boxplot( y=df["M"], x=df[a],showfliers=False )
    plt.show()
    
    chart.figure.savefig("Boxplot.png",bbox_inches='tight')
    #Generate into Document
    document.add_picture("Boxplot.png",width=Inches(6), height=Inches(3))
    boxplot_text = f"Figure 1 Boxplot for  {a} Parameter"
    paragraph = document.add_paragraph(boxplot_text)


    #Generate Statistics
    dft = df[["M",a]] 
    df_by_m = dft.groupby('M')
    df_by_m.describe().to_csv(csv_stats)

    #find Naked Eye Minimum
    dfne = df[df['M'] =='NE']
    dfne = dfne[dfne[a] == min(dfne[a])]

    #Find Optical Aided Minumum
    dfoa = df[df['M'] =='OA']
    dfoa = dfoa[dfoa[a] == min(dfoa[a])]

    #Combine Naked Eye & Optical Aided
    dfc=pd.concat([dfne, dfoa], axis=0)

    #Drop Unnesscary Column
    dfc=dfc.drop(['Moonset', 'V', 'Sunset','TZ','O',"NE",'B','T','CCD'], axis = 1)

    #Output Min Value
    #dfc.to_csv( csv_maxmin, index=False, encoding='utf-8-sig')
    
    url = csv_stats
    df = pd.read_csv(url, index_col=0)
    df=df.dropna()
    Data = ['Stats','Naked Eye', ' Optical Aided']     
    df.insert(loc=0, column='Method of Observation', value=Data)
    dfv=df.take([1,2])
    dfv.drop(columns=dfv.columns[0], axis=1,  inplace=True)
    dfv = dfv.astype(float)
    dfv = dfv.round(2)
    Data1 = ['Naked Eye', ' Optical Aided']     
    dfv.insert(loc=0, column='Method of Observation', value=Data1)
    dfz = df.take([0])
    df = pd.concat([dfz,dfv])
    
    aparameter_text = f"Table of Descriptive Statistics for  {a} Parameter"
    paragraph = document.add_paragraph(aparameter_text)


    
    t = document.add_table(df.shape[0]+1, df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
            
    aparameter_text = f"General Statistics for  {b} Parameter"
    document.add_heading(aparameter_text, level=3)
            
    #df = pd.read_csv(path, index_col=0)
    df = pd.DataFrame(data_mongo)
    df[b] = abs(df[b])
    #df=df[(df[x] <= 20)]

    #Only take Visible Data
    df = df[df['V'] =='V']

    #Only Take Evening Observation
    df = df[df['O'] =='E']

    ##Generate the Boxplot
    chart =sns.boxplot( y=df["M"], x=df[b],showfliers=False )
    plt.show()
    
    chart.figure.savefig("Boxplot.png",bbox_inches='tight')
    #Generate into Document
    document.add_picture("Boxplot.png",width=Inches(6), height=Inches(3))
    boxplot_text = f"Figure 2 Boxplot for  {b} Parameter"
    paragraph = document.add_paragraph(boxplot_text)


    #Generate Statistics
    dft = df[["M",b]] 
    df_by_m = dft.groupby('M')
    df_by_m.describe().to_csv(csv_stats)

    #find Naked Eye Minimum
    dfne = df[df['M'] =='NE']
    dfne = dfne[dfne[b] == min(dfne[b])]

    #Find Optical Aided Minumum
    dfoa = df[df['M'] =='OA']
    dfoa = dfoa[dfoa[b] == min(dfoa[b])]

    #Combine Naked Eye & Optical Aided
    dfc=pd.concat([dfne, dfoa], axis=0)

    #Drop Unnesscary Column
    dfc=dfc.drop(['Moonset', 'V', 'Sunset','TZ','O',"NE",'B','T','CCD'], axis = 1)

    #Output Min Value
    #dfc.to_csv( csv_maxmin, index=False, encoding='utf-8-sig')
    
    url = csv_stats
    df = pd.read_csv(url, index_col=0)
    df=df.dropna()
    Data = ['Stats','Naked Eye', ' Optical Aided']     
    df.insert(loc=0, column='Method of Observation', value=Data)
    dfv=df.take([1,2])
    dfv.drop(columns=dfv.columns[0], axis=1,  inplace=True)
    dfv = dfv.astype(float)
    dfv = dfv.round(2)
    Data1 = ['Naked Eye', ' Optical Aided']     
    dfv.insert(loc=0, column='Method of Observation', value=Data1)
    dfz = df.take([0])
    df = pd.concat([dfz,dfv])
    
    aparameter_text = f"Table of Descriptive Statistics for  {b} Parameter"
    paragraph = document.add_paragraph(aparameter_text)


    
    t = document.add_table(df.shape[0]+1, df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])




    #convert all varible to positive.
    #url = 'https://raw.githubusercontent.com/msyazwanfaid/hilalpy/main/Final.csv'
    paragraph = document.add_paragraph("-----------------------------------------------")
    document.add_heading('Contradiction Rate of the Criterion', level=2)
    df = pd.DataFrame(data_mongo)
    
    df = pd.DataFrame(data_mongo)
    df[a] = df[a].abs()
    df[b] = df[b].abs()

    #Set Limit

    df=df[(df[a] <= limita)]
    df=df[(df[b] <= limitb)]

    #Graph
    sns.set_theme(style="darkgrid")

    plt.figure(figsize=(10,6),dpi=1200)
    z=sns.relplot(x=df[a], y=df[b],hue=df['V'], s=20,linewidth=0.1)

    def graph(formula, x_range):
        x = np.array(x_range)
        y = eval(formula)
        plt.plot(x, y,'k', color='red')


    graph(equation, range(0, limita))

    plt.show()

    #z.savefig(figure,dpi=1200)
    z.savefig("EquaErrRate.png")

    
    #Generate into Document
    document.add_picture("EquaErrRate.png")
    figure3_text=f"Figure 3: {Criterion} Criterion Over Data of Lunar Crescent Sighting"
    paragraph = document.add_paragraph(figure3_text)


    #Condition on Whole
    print ("Total Data = ",len(df))
    
    #Generate Into Document

    paragraph = document.add_paragraph("-----------------------------------------------")
    document.add_heading('Contradiction Rate Of Whole Data', level=3)

    total_data = len(df)
    total_data_text = f"Total Data = {total_data}"
    paragraph = document.add_paragraph(total_data_text)

    x=df[a];
    df["test"]=ne.evaluate(equation)
    df_whole_pos=df[(df[b] >= df["test"])]
    df_invisible_whole_pos = df_whole_pos[df_whole_pos['V'] =='I']
    df_visible_whole_pos = df_whole_pos[df_whole_pos['V'] =='V']

    print ("Total Data Above Criteria = ",len(df_whole_pos))
    
    #Generate Into Document
    total_data_above_criteria = len(df_whole_pos)
    total_data_above_criteria_text = f"Total Data Above Criteria = {total_data_above_criteria}"
    paragraph = document.add_paragraph(total_data_above_criteria_text)

    
    positive_errorrate_whole = ((len(df_invisible_whole_pos))/(len(df_whole_pos)))*100

    print ("Total Positive Contradiction = ",len(df_invisible_whole_pos))
    
    #Generate Into Document
    total_positive_contradiction = len(df_invisible_whole_pos)
    total_positive_contradiction_text = f"Total Positive Contradiction = {total_positive_contradiction}"
    paragraph = document.add_paragraph(total_positive_contradiction_text)



    df_whole_neg=df[(df[b] <= df["test"])]
    df_visible_whole_neg =  df_whole_neg[df_whole_neg['V'] =='V']
    df_invisible_whole_neg = df_whole_neg[df_whole_neg['V'] =='I']

    negative_errorrate_whole = abs((len(df_visible_whole_neg))/(len(df_whole_neg)))*100
    
    print ("Total Data Below Criteria = ",len(df_whole_neg))
    #Generate Into Document
    total_data_below_criteria = len(df_whole_neg)
    total_data_below_criteria_text = f"Total Data Below Criteria = {total_data_below_criteria}"
    paragraph = document.add_paragraph(total_data_below_criteria_text)

    print ("Total Negative Contradiction = ",len(df_visible_whole_neg))
    #Generate Into Document
    total_negative_contradiction = len(df_visible_whole_neg)
    total_negative_contradiction_text = f"Total Negative Contradiction = {total_negative_contradiction}"
    paragraph = document.add_paragraph(total_negative_contradiction_text)
    print("")
    
    #Combine Dataframe
    df_positiveerrorrate = df_invisible_whole_pos
    df_negativeerrorrate = df_visible_whole_neg
    df_invisible_whole_pos.to_csv( poserrorratedata, index=False, encoding='utf-8-sig')
    df_visible_whole_neg.to_csv( negerrorratedata, index=False, encoding='utf-8-sig')


    #Condition Test on Naked Eye
    dfn = df[df['M'] =='NE']
    df_ne_pos=dfn[(dfn[b] >= dfn["test"])]
    df_ne_pos_invisible = df_ne_pos[df_ne_pos['V'] =='I']
    df_ne_pos_visible = df_ne_pos[df_ne_pos['V'] =='V']

    positive_errorrate_nakedeye = ((len(df_ne_pos_invisible))/(len(df_ne_pos)))*100
        
    #Generate Into Document
    paragraph = document.add_paragraph("-----------------------------------------------")
    
    print ("Total Data Above Criteria (NE) = ",len(df_ne_pos))
    document.add_heading('Contradiction Rate Of Naked Eye (NE) Data', level=3)
    total_data_above_criteria_NE = len(df_ne_pos)
    total_data_above_criteria_NE_text = f"Total Data Above Criteria (NE) = {total_data_above_criteria_NE}"
    paragraph = document.add_paragraph(total_data_above_criteria_NE_text)

    #Generate Into Document
    print ("Total Positive Contradiction (NE) = ",len(df_ne_pos_invisible))
    total_positive_contradiction_NE = len(df_ne_pos_invisible)
    total_positive_contradiction_NE_text = f"Total Positive Contradiction (NE) = {total_positive_contradiction_NE}"
    paragraph = document.add_paragraph(total_positive_contradiction_NE_text)


    df_ne_neg=dfn[(dfn[b] <= dfn["test"])]
    df_ne_neg_visible = df_ne_neg[df_ne_neg['V'] =='V']
    df_ne_neg_invisible = df_ne_neg[df_ne_neg['V'] =='I']

    negative_errorrate_nakedeye = ((len(df_ne_neg_visible))/(len(df_ne_neg)))*100
    
    print ("Total Data Below Criteria (NE) = ",len(df_ne_neg))
    #Generate Into Document
    Total_Data_Below_Criteria_NE = len(df_ne_neg)
    Total_Data_Below_Criteria_NE_text = f"Total Data Below Criteria (NE) = {Total_Data_Below_Criteria_NE}"
    paragraph = document.add_paragraph(Total_Data_Below_Criteria_NE_text)
    
    print ("Total Negative Contradiction  (NE) = ",len(df_ne_neg_visible))
    #Generate Into Document
    Total_Negative_Contradiction_NE = len(df_ne_neg_visible)
    Total_Negative_Contradiction_NE_text = f"Total Negative Contradiction  (NE) = {Total_Negative_Contradiction_NE}"
    paragraph = document.add_paragraph(Total_Negative_Contradiction_NE_text)
    
    
    print("")
    paragraph = document.add_paragraph("-----------------------------------------------")


    #Condition on Optical Aided
    dfb = df[df['M'] =='OA']

    df_oa_pos=dfb[(dfb[b] >= dfb["test"])]
    df_oa_pos_invisible = df_oa_pos[df_oa_pos['V'] =='I']
    df_oa_pos_visible = df_oa_pos[df_oa_pos['V'] =='V']

    positive_errorrate_opticalaided  = ((len(df_oa_pos_invisible))/(len(df_oa_pos)))*100
    
    #Generate Into Document
    document.add_heading('Contradiction Rate Of Optical Aided (OA) Data', level=3)
    print ("Total Data Above Criteria (OA) = ",len(df_oa_pos))
    Total_Data_Above_Criteria_OA = len(df_oa_pos)
    Total_Data_Above_Criteria_OA_text = f"Total Data Above Criteria (OA) = {Total_Data_Above_Criteria_OA}"
    paragraph = document.add_paragraph(Total_Data_Above_Criteria_OA_text)
    
    
    
    #Generate Into Document
    print ("Total Positive Contradiction (OA) = ",len(df_oa_pos_invisible))
    Total_Positive_Contradiction_OA = len(df_oa_pos_invisible)
    Total_Positive_Contradiction_OA_text = f"Total Positive Contradiction (OA) = {Total_Positive_Contradiction_OA}"
    paragraph = document.add_paragraph(Total_Positive_Contradiction_OA_text)
    
    

    df_oa_neg=dfb[(dfb[b] <= dfb["test"])]
    df_oa_neg_visible = df_oa_neg[df_oa_neg['V'] =='V']
    df_oa_neg_invisible = df_oa_neg[df_oa_neg['V'] =='I']

    try:
        negative_errorrate_opticalaided = ((len(df_oa_neg_visible))/(len(df_oa_neg)))*100
    except ZeroDivisionError as err:
        negative_errorrate_opticalaided = 0
    
    #Generate Into Document
    print ("Total Data Below Criteria (OA) = ",len(df_oa_neg))
    Total_Data_Below_Criteria_OA = len(df_oa_neg)
    Total_Data_Below_Criteria_OA_text = f"Total Data Below Criteria (OA) = {Total_Data_Below_Criteria_OA}"
    paragraph = document.add_paragraph(Total_Data_Below_Criteria_OA_text)
    
    
    
    #Generate Into Document
    print ("Total Negative Contradiction  (OA) = ",len(df_oa_neg_visible))
    Total_Negative_Contradiction_OA = len(df_oa_neg_visible)
    Total_Negative_Contradiction_OA_text = f"Total Negative Contradiction  (OA) = {Total_Negative_Contradiction_OA}"
    paragraph = document.add_paragraph(Total_Negative_Contradiction_OA_text)
    
    print("")
    paragraph = document.add_paragraph("-----------------------------------------------")
    document.add_heading('Contradiction Rate Summary in Percentage', level=3)


        
    #Merge Error Rate
    #df = pd.merge(dfy_visible, df_visible, how='outer', indicator=True).query("_merge != 'both'").drop('_merge', axis=1).reset_index(drop=True)
    #dfccd = df[df['I'] =='CCD']
    #dfNU = df[df['I'] =='NU']
    #dfT = df[df['I'] =='T']

    condition_test_result = {'Parameter': ['Whole','Naked Eye','Optical Aided'],
            'Positive': [positive_errorrate_whole,positive_errorrate_nakedeye,positive_errorrate_opticalaided],
            'Negative': [negative_errorrate_whole,negative_errorrate_nakedeye,negative_errorrate_opticalaided]
                            }
    
    df_cond_result = pd.DataFrame(condition_test_result, columns = ['Parameter', 'Positive','Negative'])
    df=df_cond_result.round(2)
    
    
    #Document Generation
    t = document.add_table(df.shape[0]+1, df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])

    paragraph = document.add_paragraph("-----------------------------------------------")


    print (df_cond_result)

    #df.to_csv( errorratetotal, index=False, encoding='utf-8-sig')

    #Whole General Ranking Determination
    heading_text =f'Comparison of the {Criterion} Criterion Against other Criterion'
    document.add_heading('Comparison of the Tested Criterion Against other Criterion', level=2)
    paragraph = document.add_paragraph("Ranking for general contradiction is based on the mean percentage value of both positive and negative contradiction of a criterion")

    document.add_heading('Whole Contradiction Percentage', level=3)

    client = MongoClient("mongodb+srv://userhilalpy:kSTkJBzJ8CbCATmm@hilalpy.wdazfdi.mongodb.net/?retryWrites=true&w=majority")
    db = client.test
    db = client['datahillalpy']
    collection = db['whole_csv']
    
    data_mongo_whole = []
    for doc in collection.find():
        data_mongo_whole.append(doc)
    
    df_whole = pd.DataFrame(data_mongo_whole)
    df1 = pd.DataFrame({
        "Contradiction": ["Positive", "Negative"],
        "Data": ["Tested Data","Tested Data"],
        "Percentage" : [positive_errorrate_whole,negative_errorrate_whole]
        })
    df_whole_new=pd.concat([df_whole, df1])
    df_whole_new=df_whole_new.round(2).sort_values('Data')
    
    #Whole Bar Plot
    plt.figure(figsize=(12, 6))
    chart = sns.barplot(x=df_whole_new.Data, y=df_whole_new.Percentage,
                        hue=df_whole_new.Contradiction)
    chart.set_xticklabels(chart.get_xticklabels(), rotation=90, horizontalalignment='right')
    plt.show()
    chart.figure.savefig("WholeBarPlot.png",bbox_inches='tight')
    
    #Generate into Document
    document.add_picture("WholeBarPlot.png",width=Inches(6), height=Inches(3))
    paragraph = document.add_paragraph("Figure 4: Comparison of Whole Contradiction Rate for Tested Criterion against other Criterion")


    #Ranking for Whole Contradiction

    result_df = pd.DataFrame(columns=['Result_1', 'Result_2'])
    df_whole_ranking = df_whole_new[['Percentage', 'Data']]
    
    y = df_whole_ranking.Percentage.values
    v = pd.DataFrame({'Data' : df_whole_ranking.Data.values[::2], 'Percentage' : y[::2] + y[1::2]})    
    v['Percentage'] = v['Percentage'].div(2).round(2)   
    df_whole_ranked = v.sort_values('Percentage')
    df_whole_ranked.insert(0, 'Ranking', range(1, 1 + len(df_whole_ranked)))
    df=df_whole_ranked
    
    df_whole_new = df_whole_new.pivot(index='Data', columns='Contradiction', values='Percentage')
    df_whole_new.columns = ['Negative Contradiction', 'Positive Contradiction']
    df = df.merge(df_whole_new, on='Data', how='outer')
    df = df.rename(columns={'Percentage': 'Mean Percentage'})


    
    t = document.add_table(df.shape[0]+1, df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
            
    #Naked Eye General Ranking Determination

    document.add_heading('Naked Eye General Contradiction Percentage', level=3)
    
    
    client = MongoClient("mongodb+srv://userhilalpy:kSTkJBzJ8CbCATmm@hilalpy.wdazfdi.mongodb.net/?retryWrites=true&w=majority")
    db = client.test
    db = client['datahillalpy']
    collection = db['nakedeye_csv']
    
    data_mongo_ne = []
    for doc in collection.find():
        data_mongo_ne.append(doc)
    
    df_calendar = pd.DataFrame(data_mongo_ne)
    #df_calendar = pd.read_csv(path2)
    df1 = pd.DataFrame({
        "Contradiction": ["Positive", "Negative"],
        "Data": ["Tested Data","Tested Data"],
        "Percentage" : [positive_errorrate_nakedeye,negative_errorrate_nakedeye]
        })
    df_general_ne=pd.concat([df_calendar, df1])
    df_general_ne=df_general_ne.round(2).sort_values('Data')
    df_general_ne_rnk=df_general_ne.round(2).sort_values('Data')
        
        #Naked Eye Bar Plot
    plt.figure(figsize=(12, 6))
    chart = sns.barplot(x=df_general_ne.Data, y=df_general_ne.Percentage,
                        hue=df_general_ne.Contradiction)
    chart.set_xticklabels(chart.get_xticklabels(), rotation=90, horizontalalignment='right')
    plt.show()
    chart.figure.savefig("NakedEyeBarPlot.png",bbox_inches='tight')
    
        #Generate into Document
    document.add_picture("NakedEyeBarPlot.png",width=Inches(6), height=Inches(3))
    paragraph = document.add_paragraph("Figure 5: Comparison of Naked Eye Contradiction Rate for Tested Criterion against other Criterion")
    
        #Naked Eye General Ranking Determination
    
    result_df = pd.DataFrame(columns=['Result_1', 'Result_2'])
    df_generalne_ranking = df_general_ne[['Percentage', 'Data']]
    
    y = df_generalne_ranking.Percentage.values
    v = pd.DataFrame({'Data' : df_generalne_ranking.Data.values[::2], 'Percentage' : y[::2] + y[1::2]})    
    v['Percentage'] = v['Percentage'].div(2).round(2)   
    df_ne_ranked = v.sort_values('Percentage')
    df_ne_ranked.insert(0, 'Ranking', range(1, 1 + len(df_ne_ranked)))
    df=df_ne_ranked
    
    df_general_ne2 = df_general_ne_rnk.pivot(index='Data', columns='Contradiction', values='Percentage')
    df_general_ne2.columns = ['Negative Contradiction', 'Positive Contradiction']
    df = df.merge(df_general_ne2, on='Data', how='outer')
    df = df.rename(columns={'Percentage': 'Mean Percentage'})
        
    
    t = document.add_table(df.shape[0]+1, df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
            
    #Optical Aided General Ranking Determination

    document.add_heading('Optical Aided General Contradiction Percentage', level=3)
    
    client = MongoClient("mongodb+srv://userhilalpy:kSTkJBzJ8CbCATmm@hilalpy.wdazfdi.mongodb.net/?retryWrites=true&w=majority")
    db = client.test
    db = client['datahillalpy']
    collection = db['opticalaided_csv']
    
    data_mongo_oa = []
    for doc in collection.find():
        data_mongo_oa.append(doc)
    
    df_oa = pd.DataFrame(data_mongo_oa)
    df1 = pd.DataFrame({
        "Contradiction": ["Positive", "Negative"],
        "Data": ["Tested Data","Tested Data"],
        "Percentage" : [positive_errorrate_opticalaided,negative_errorrate_opticalaided]
        })
    df_general_oa=pd.concat([df_oa, df1])
    df_general_oa=df_general_oa.round(2).sort_values('Data')
    df_general_oa_rnk=df_general_oa.sort_values('Data')


        #Optical Aided Bar Plot
    plt.figure(figsize=(12, 6))
    chart = sns.barplot(x=df_general_oa.Data, y=df_general_oa.Percentage,
                        hue=df_general_oa.Contradiction)
    chart.set_xticklabels(chart.get_xticklabels(), rotation=90, horizontalalignment='right')
    plt.show()
    chart.figure.savefig("OpticalAidedBarPlot.png",bbox_inches='tight')
    document.add_picture("OpticalAidedBarPlot.png",width=Inches(6), height=Inches(3))
    paragraph = document.add_paragraph("Figure 6: Comparison of Optical Aided Contradiction Rate for Tested Criterion against other Criterion")

        #Optical Aided General Ranking Determination
    result_df = pd.DataFrame(columns=['Result_1', 'Result_2'])
    df_generaloa_ranking = df_general_oa[['Percentage', 'Data']]
    
    y = df_generaloa_ranking.Percentage.values
    v = pd.DataFrame({'Data' : df_generaloa_ranking.Data.values[::2], 'Percentage' : y[::2] + y[1::2]})    
    v['Percentage'] = v['Percentage'].div(2).round(2)   
    df_oa_ranked = v.sort_values('Percentage')
    df_oa_ranked.insert(0, 'Ranking', range(1, 1 + len(df_oa_ranked)))
    df=df_oa_ranked
    
    df_general_oa2 = df_general_oa_rnk.pivot(index='Data', columns='Contradiction', values='Percentage')
    df_general_oa2.columns = ['Negative Contradiction', 'Positive Contradiction']
    df = df.merge(df_general_oa2, on='Data', how='outer')
    df = df.rename(columns={'Percentage': 'Mean Percentage'})
    
    t = document.add_table(df.shape[0]+1, df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
   

    #Calendrical Purpose Ranking Naked Eye
    document.add_heading('Comparison of the Tested Criterion Against other Criterion (For Calendrical Purposes)', level=2)
    paragraph = document.add_paragraph("Ranking for calendrical based criterion is based on percentage of negative contradiction. The lower is better.")
    document.add_heading('Calendrical Purposed for Naked Eye', level=3)


    
    #Ranking for Naked Eye Calendrical Contradiction


    result_df = pd.DataFrame(columns=['Result_1', 'Result_2'])
    df_general_ne1 = df_general_ne[df_general_ne["Contradiction"].str.contains("Positive") == False]
    df_calendar_ne = df_general_ne1[['Percentage', 'Data']]
    
     
    df_calendarne_ranked = df_calendar_ne.sort_values('Percentage')
    df_calendarne_ranked.insert(0, 'Ranking', range(1, 1 + len(df_calendarne_ranked)))
    df=df_calendarne_ranked
    t = document.add_table(df.shape[0]+1, df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
            
    paragraph = document.add_paragraph("-----------------------------------------------")

            
    #Calendrical Purpose Ranking Optical Aided
    document.add_heading('Calendrical Purposed for Optical Aided', level=3)
  
        
    #Ranking for Optical Aided Calendrical Contradiction

    result_df = pd.DataFrame(columns=['Result_1', 'Result_2'])
    df_general_oa1 = df_general_oa[df_general_oa["Contradiction"].str.contains("Positive") == False]
    df_calendar_oa = df_general_oa1[['Percentage', 'Data']]
    
     
    df_calendaroa_ranked = df_calendar_oa.sort_values('Percentage')
    df_calendaroa_ranked.insert(0, 'Ranking', range(1, 1 + len(df_calendaroa_ranked)))
    df=df_calendaroa_ranked
    t = document.add_table(df.shape[0]+1, df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
            
    paragraph = document.add_paragraph("-----------------------------------------------")
    
    #Ranking for Observation Purposed
    document.add_heading('Comparison of the Tested Criterion Against other Criterion (Observation Purposes)', level=2)
    paragraph = document.add_paragraph("Ranking for  criterion is based on percentage of positive contradiction. The lower is better.")
    document.add_heading('Observation Purposed for Naked Eye', level=3)
    
        #Ranking for Naked Eye Calendrical Contradiction


    result_df = pd.DataFrame(columns=['Result_1', 'Result_2'])
    df_general_ne = df_general_ne[df_general_ne["Contradiction"].str.contains("Negative") == False]
    df_calendar_ne = df_general_ne[['Percentage', 'Data']]
    
     
    df_calendarne_ranked = df_calendar_ne.sort_values('Percentage')
    df_calendarne_ranked.insert(0, 'Ranking', range(1, 1 + len(df_calendarne_ranked)))
    df=df_calendarne_ranked
    t = document.add_table(df.shape[0]+1, df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
            
    document.add_heading('Observation Purposed for Optical Aided', level=3)
    
    result_df = pd.DataFrame(columns=['Result_1', 'Result_2'])
    df_general_oa = df_general_oa[df_general_oa["Contradiction"].str.contains("Negative") == False]
    df_calendar_oa = df_general_oa[['Percentage', 'Data']]
    
     
    df_calendaroa_ranked = df_calendar_oa.sort_values('Percentage')
    df_calendaroa_ranked.insert(0, 'Ranking', range(1, 1 + len(df_calendaroa_ranked)))
    df=df_calendaroa_ranked
    t = document.add_table(df.shape[0]+1, df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
    
    #Regression Assesment        
    document.add_heading('Criterion Regression Assesment', level=2)
    
        #Resorted the Data

    
    df = pd.DataFrame(data_mongo)
    df[a] = df[a].abs()
    df[b] = df[b].abs()
    #print(dataset)

        
    dftest = df[df['V'] =='V']
    #if dataset == "Whole":
    #    dftest=dftest
    #else:
    #    dftest = dftest[dftest['M'] ==dataset]

    dftest = dftest[[a,b,'V','M']]

    df_test_sorted = dftest.sort_values(a)
    df_test_sorted = df_test_sorted[df_test_sorted[a] <= limita]
    df_test_sorted = df_test_sorted[df_test_sorted[b] <= limitb]
    
            # Create a Selected DataFrame
    df = df_test_sorted
    min_range = 1
    max_range = limita
    increment = 1
    results = []

    for i in range(min_range, max_range + 1, increment):
        mask = (df[a] >= i) & (df[a] < i + increment)
        subset = df[mask]
        if len(subset) > 0:
            min_b = subset[b].min()
            results.append({'x': f'{i + increment - 1}', 'y': min_b})

    result_df = pd.DataFrame(results)
    df = result_df
    df['x'] = df['x'].astype(np.int64)
    

        # Calculate the accuracy metrics
        
    x = df['x'].astype(float).values.reshape(-1, 1)
    y_pred = eval(equation)
        
    mae = mean_absolute_error(df['y'], y_pred.ravel())
    mse = mean_squared_error(df['y'], y_pred.ravel())
    r2 = r2_score(df['y'], y_pred.ravel())
    
    MAE_text = f"Mean Absolute Error (MAE) = {mae}"
    paragraph = document.add_paragraph(MAE_text)
    MSE_text = f"Mean Squared Error (MSE) = {mse}"
    paragraph = document.add_paragraph(MSE_text)
    R2_text = f"R^2 Score = {r2}"
    paragraph = document.add_paragraph(R2_text)

        # Plot the actual values against the predicted values
    plt.figure(figsize=(20,10))
    chart=plt.scatter(df['x'], df['y'], color='blue', label='Actual')
    plt.plot(x, y_pred, color='red', label='Predicted')
    plt.xlabel(a)
    plt.ylabel(b)
    plt.legend()
    plt.show()
    chart.figure.savefig("RegressionAssesment.png",bbox_inches='tight')
    document.add_picture("RegressionAssesment.png",width=Inches(6), height=Inches(3))
    paragraph = document.add_paragraph("Figure 7: Graph of the Criterion over Minimum Value of Lunar Crescent Data")

    client = MongoClient("mongodb+srv://userhilalpy:kSTkJBzJ8CbCATmm@hilalpy.wdazfdi.mongodb.net/?retryWrites=true&w=majority")
    db = client.test
    db = client['datahillalpy']
    collection = db['reg_exam']
    
    data_mongo_reg_ex = []
    for doc in collection.find():
        data_mongo_reg_ex.append(doc)
    
    df = pd.DataFrame(data_mongo_reg_ex)
    
    df1 = pd.DataFrame({
    "Mean Absolute Error (MAE)": [mae],
    "Mean Squared Error (MSE)": [mse],
    "R^2 Score" : [r2],
    "Criterion" : ["Tested Data"],
    })
    df=pd.concat([df, df1])
    df.drop(columns='_id', inplace=True)
    df=df.round(2)
    #print(df)
    #print(dataset)
    #df['mean'] = df.mean(axis=1)
    df.sort_values(by='Mean Absolute Error (MAE)', inplace=True)
    df['Rank'] = range(1, len(df) + 1)
    #df.drop(columns='mean', inplace=True)

    
    document.add_heading('Criterion Ranking based on Regression', level=3)

    
    t = document.add_table(df.shape[0]+1, df.shape[1])
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])

    
    #chart.figure.savefig("OpticalAidedBarPlot.png",bbox_inches='tight')
    #document.add_picture("OpticalAidedBarPlot.png",width=Inches(6), height=Inches(3))

    document.save(doc_loc)



